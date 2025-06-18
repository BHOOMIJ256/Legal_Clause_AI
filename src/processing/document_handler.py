from typing import Dict, List, Optional
import PyPDF2
import docx
import io

class DocumentHandler:
    def __init__(self):
        """Initialize the DocumentHandler with supported file types."""
        self.supported_types = ['pdf', 'docx', 'txt']
    
    def extract_text_from_pdf(self, pdf_file: bytes) -> str:
        """
        Extract text from a PDF file in memory.
        
        Args:
            pdf_file (bytes): PDF file content as bytes
            
        Returns:
            str: Extracted text from the PDF
        """
        try:
            print("Extracting text from PDF...")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file))
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                print(f"Processed page {i+1}")
            
            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")
                
            print(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_text_from_docx(self, docx_file: bytes) -> str:
        """
        Extract text from a DOCX file in memory.
        
        Args:
            docx_file (bytes): DOCX file content as bytes
            
        Returns:
            str: Extracted text from the DOCX
        """
        try:
            print("Extracting text from DOCX...")
            doc = docx.Document(io.BytesIO(docx_file))
            text = ""
            for i, para in enumerate(doc.paragraphs):
                if para.text:
                    text += para.text + "\n"
                print(f"Processed paragraph {i+1}")
            
            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX")
                
            print(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_text_from_txt(self, txt_file: bytes) -> str:
        """
        Extract text from a TXT file in memory.
        
        Args:
            txt_file (bytes): TXT file content as bytes
            
        Returns:
            str: Extracted text from the TXT
        """
        try:
            print("Extracting text from TXT...")
            text = txt_file.decode('utf-8')
            
            if not text.strip():
                raise ValueError("No text could be extracted from the TXT")
                
            print(f"Successfully extracted {len(text)} characters from TXT")
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from TXT: {str(e)}")
            raise
    
    def process_document(self, file_content: bytes, file_type: str) -> str:
        """
        Process a document and extract text.
        
        Args:
            file_content (bytes): Raw file content
            file_type (str): Type of file ('pdf', 'docx', 'txt')
            
        Returns:
            str: Extracted text content
        """
        if not isinstance(file_content, bytes):
            raise ValueError(f"File content must be bytes, got {type(file_content)}")
            
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        print(f"\nProcessing {file_type.upper()} document...")
        
        try:
            if file_type == 'pdf':
                text = self._extract_from_pdf(file_content)
            elif file_type == 'docx':
                text = self._extract_from_docx(file_content)
            elif file_type == 'txt':
                text = self._extract_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
            # Ensure we have a string
            if not isinstance(text, str):
                text = str(text)
                
            # Clean and normalize the text
            text = text.strip()
            
            if not text:
                raise ValueError("No text content extracted from document")
                
            print(f"Successfully extracted {len(text)} characters from {file_type.upper()}")
            return text
            
        except Exception as e:
            print(f"Error processing {file_type} document: {str(e)}")
            raise

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            raise

    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            text = file_content.decode('utf-8')
            return text.strip()
        except UnicodeDecodeError:
            # Try with a different encoding if UTF-8 fails
            try:
                text = file_content.decode('latin-1')
                return text.strip()
            except Exception as e:
                print(f"Error decoding text file: {str(e)}")
                raise
        except Exception as e:
            print(f"Error extracting text from TXT: {str(e)}")
            raise 